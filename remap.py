import pandas as pd
from tkinter import Tk, filedialog, messagebox, Button, Label, Toplevel
from docx import Document
from docx.shared import Pt
import os
import re
import sys
import os

prefixes_to_names = {
    "C": "Конденсаторы", "R": "Резисторы", "D": "Диоды", "G": "Генераторы", "F": "Предохранители", "L": "Индуктивные элементы", "VD": "Схема интегральная", "VT": "Транзисторы", "T": "Трансформаторы", "SB": "Выключатель кнопочный", "X": "Соединения контактные", "Z": "Фильтры", "AB": "Приводы исполнительных механизмов", "АС": "Устройство АВР", "RC": "Чип-резисторы",
    "AF": "Регулятор частоты", "АК": "Устройство (комплект) реле защит", "АКB": "Устройство блокировки типа КРБ", "AKS": "Устройство АПВ", "AKV": "Устройство комплектное продольной дифзащиты ЛЭП", "AKZ": "Устройство комплектное реле сопротивления",
    "AR": "Устройство комплектное реле УРОВ", "AV": "Устройство регулирования напряжения", "AW": "Регулятор мощности", "ВА": "Громкоговоритель", "ВВ": "Магнитострикционный элемент", "BD": "Детектор ионизирующих излучений", 
    "BE": "Сельсин-приемник", "ВF": "Телефон (капсюль)", "ВС": "Сельсин-датчик", "ВК": "Тепловой датчик", "BL": "Фотоэлемент", "ВМ": "Микрофон", "ВР": "Датчик давления", "BQ": "Пьезоэлемент", "BR": "Датчик частоты вращения",
    "BS": "Звукосниматель", "BV": "Датчик скорости", "BT": "Датчик температуры", "BVA": "Счетчик вольтамперчасов реактивных", "BW": "Счетчик ватт-часов активных", "CB": "Конденсаторный силовой блок", "CG": "Конденсаторный зарядный блок",
    "DA": "Схема интегральная аналоговая", "DD": "Схема интегральная цифровая", "DS": "Устройства хранения информации", "DT": "Устройство задержки", "ЕК": "Нагревательный элемент", "EL": "Лампа осветительная", "ET": "Пиропатрон",
    "FA": "Дискретный элемент защиты по току мгновенного действия", "FP": "Дискретный элемент защиты по току инерционного действия", "FU": "Предохранитель плавкий", "FV": "Дискретный элемент защиты по напряжению, разрядник", "GB": "Батарея",
    "GC": "Синхронный компенсаторы", "GE": "Возбудитель генератора", "GEA": "Подвозбудитель (вспомогательный возбудитель)", "НА": "Прибор звуковой сигнализации", "HG": "Индикатор символьный", "HL": "Приборы световой сигнализации",
    "HLA": "Световое табло", "LG": "Лампа сигнализации с линзой зеленой", "HLR": "Лампа сигнализации с линзой красной", "HLW": "Лампа сигнализации с линзой белой", "HY": "Индикатор полупроводниковый", "КА": "Реле токовое", 
    "КН": "Реле указательное", "КК": "Реле электротепловое", "КМ": "Контактор, магнитный пускатель", "КТ": "Реле времени", "KV": "Реле напряжения", "KA0": "Реле тока нулевой последовательности", "KAT": "Реле тока с насыщающимся трансформатором",
    "KAW": "Реле тока с торможением", "KAZ": "Реле тока фильтровое", "KB": "Реле блокировки", "KBS": "Реле блокировки от многократных включений", "KCC": "Реле команды «включить»", "KCT": "Реле команды «отключить»", "KF": "Реле частоты", "KHA": "Реле импульсной сигнализации",
    "KLP": "Реле давления повторительное", "KQ": "Реле фиксации положения выключателя", "KQC": "Реле положения «Включено»", "KQT": "Реле положения «Отключено»", "KQQ": "Реле фиксации команды включения", "KQS": "Реле фиксации положения разъединителя",
    "KS": "Реле контроля", "KSG": "Реле газовое", "KSH": "Реле струи (напора)", "KSS": "Реле контроля синхронизма", "KSV": "Реле контроля напряжения", "KZ": "Реле сопротивления", "KVZ": "Фильтр – реле напряжения", "KW": "Реле мощности"
}

def resource_path(relative_path):
    """Получает абсолютный путь к ресурсу."""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

def process_and_fill_template(file_path, template_path, loading_screen):
    try:
        data = pd.read_excel(file_path, engine='openpyxl')

        df = pd.DataFrame({
            "Поз. обозначение": data['Part Reference'].astype(str),
            "Наименование": data['Manufacturer Part Number'].astype(str) + ' ' + data['Manufacturer'].astype(str),
            "Кол.": data['Quantity'].astype(int),
            "Примечание": data['Value'].astype(str) + ' ±' + data['Tolerance'].astype(str)
        })

        df = df.sort_values(by="Поз. обозначение", key=lambda col: col.str.extract(r"([A-Z]+)(\d+)").apply(lambda x: (x[0], int(x[1])), axis=1))

        rows_with_names = []
        temp_designations = []
        temp_rows = []
        added_prefixes = set()

        def process_group(designations, rows):
            processed_rows = []
            numbers = [int(re.search(r"\d+", d).group()) for d in designations]
            if len(designations) > 1 and any(num > 99 for num in numbers):
                designation_str = f"{designations[0]}-"
                processed_rows.append({"Поз. обозначение": designation_str, "Наименование": rows[0]["Наименование"], "Кол.": "", "Примечание": ""})
                processed_rows.append({"Поз. обозначение": designations[-1], "Наименование": "", "Кол.": sum(row["Кол."] for row in rows), "Примечание": rows[0]["Примечание"]})
            else:
                designation_str = f"{designations[0]}-{designations[-1]}" if len(designations) > 2 else ",".join(designations)
                processed_rows.append({"Поз. обозначение": designation_str, "Наименование": rows[-1]["Наименование"], "Кол.": sum(row["Кол."] for row in rows), "Примечание": rows[0]["Примечание"]})
            return processed_rows

        for _, row in df.iterrows():
            prefix = re.match(r"([A-Za-z]+)", row["Поз. обозначение"]).group(1)

            if temp_designations and (temp_rows[-1]["Наименование"] != row["Наименование"] or re.match(r"([A-Za-z]+)", temp_designations[-1]).group(1) != prefix):
                rows_with_names += process_group(temp_designations, temp_rows)
                temp_designations = []
                temp_rows = []

            if prefix not in added_prefixes:
                added_prefixes.add(prefix)
                rows_with_names.append({"Поз. обозначение": " ", "Наименование": "", "Кол.": "", "Примечание": ""})
                rows_with_names.append({"Поз. обозначение": " ", "Наименование": prefixes_to_names.get(prefix, "Неизвестное наименование"), "Кол.": "", "Примечание": ""})

            temp_designations.append(row["Поз. обозначение"])
            temp_rows.append(row)

        if temp_designations:
            rows_with_names += process_group(temp_designations, temp_rows)

        doc = Document(template_path)
        table = doc.tables[0]

        font_name = 'GOST_A.TTF'
        font_size = Pt(12)

        start_row_index = 1
        rows_per_page = 30  # Примерное количество строк на странице

        for i, row in enumerate(rows_with_names, start=start_row_index):
            if i % rows_per_page == 0 and i > start_row_index:
                if i % rows_per_page == 0 and i > start_row_index:
                    doc.add_paragraph("")
                    doc.add_paragraph("", style = None).paragraph_format.page_break_before

            if i >= len(table.rows):
                table.add_row()
            row_cells = table.rows[i].cells

            row_cells[0].text = row['Поз. обозначение']
            row_cells[1].text = row['Наименование']
            row_cells[2].text = str(row['Кол.'])
            row_cells[3].text = row['Примечание']

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = font_size
                        run.italic = True

        filled_row_count = start_row_index + len(rows_with_names)
        total_row_count = len(table.rows)

        for row_index in range(total_row_count - 1, filled_row_count - 1, -1):
            table._tbl.remove(table.rows[row_index]._tr)

        output_file = os.path.splitext(file_path)[0] + "_filled_template.docx"
        doc.save(output_file)

        loading_screen.destroy()
        messagebox.showinfo("Успех", f"Документ Word успешно создан: {output_file}")

    except Exception as e:
        loading_screen.destroy()
        messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")

def select_file():
    file_path = filedialog.askopenfilename(title="Выберите Excel файл", filetypes=[("Excel files", "*.xlsx")])
    if not file_path:
        return

    template_path = os.path.join(os.path.dirname(__file__), "template.docx")
    if not os.path.exists(template_path):
        messagebox.showerror("Ошибка", "Шаблон Word (template.docx) не найден в папке с программой.")
        return

    loading_screen = Toplevel(root)
    loading_screen.title("Загрузка")
    Label(loading_screen, text="Подождите, идет обработка данных...").pack(padx=20, pady=20)
    loading_screen.geometry("300x100")
    loading_screen.update()

    root.after(100, lambda: process_and_fill_template(file_path, template_path, loading_screen))

root = Tk()
root.title("Конвертер Excel в Word (группировка)")

Label(root, text="Выберите Excel файл для конвертации").pack(pady=10)
Button(root, text="Выбрать файл", command=lambda: select_file()).pack(pady=10)

root.mainloop()

